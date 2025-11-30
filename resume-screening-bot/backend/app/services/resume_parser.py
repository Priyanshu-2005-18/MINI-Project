import PyPDF2
import pdfplumber
from docx import Document
from pathlib import Path
import re

class ResumeParser:
    """Parse resumes from PDF, DOCX, and text files"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text
    
    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        """Extract text from various file formats"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == ".pdf":
            return ResumeParser.extract_text_from_pdf(file_path)
        elif file_extension in [".docx", ".doc"]:
            return ResumeParser.extract_text_from_docx(file_path)
        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def parse_resume_structure(text: str) -> dict:
        """Parse resume into standardized structure with enhanced extraction"""
        parsed = {
            "personal_info": {},
            "education": [],
            "experience": [],
            "technical_skills": [],
            "projects": [],
            "certifications": [],
            "achievements": []
        }
        
        # Email extraction
        emails = re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text)
        if emails:
            parsed["personal_info"]["email"] = emails[0]
        
        # Phone extraction
        phones = re.findall(r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b", text)
        if phones:
            parsed["personal_info"]["phone"] = phones[0]
        
        # Enhanced section detection
        lines = text.split("\n")
        current_section = None
        section_content = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            line_lower = line_stripped.lower()
            
            # Detect section headers
            if re.match(r"^(education|academic|qualification)", line_lower):
                if section_content and current_section:
                    parsed[current_section].extend(section_content)
                current_section = "education"
                section_content = []
            elif re.match(r"^(experience|work|employment|professional|career)", line_lower):
                if section_content and current_section:
                    parsed[current_section].extend(section_content)
                current_section = "experience"
                section_content = []
            elif re.match(r"^(skill|technical|competenc)", line_lower):
                if section_content and current_section:
                    parsed[current_section].extend(section_content)
                current_section = "technical_skills"
                section_content = []
            elif re.match(r"^(project|portfolio|work sample)", line_lower):
                if section_content and current_section:
                    parsed[current_section].extend(section_content)
                current_section = "projects"
                section_content = []
            elif re.match(r"^(certification|certificate)", line_lower):
                if section_content and current_section:
                    parsed[current_section].extend(section_content)
                current_section = "certifications"
                section_content = []
            elif re.match(r"^(achievement|award|honor)", line_lower):
                if section_content and current_section:
                    parsed[current_section].extend(section_content)
                current_section = "achievements"
                section_content = []
            
            # Add content to current section
            if current_section:
                if current_section == "technical_skills":
                    # Split skills by common delimiters
                    skills = re.split(r"[,;|•\-\*]", line_stripped)
                    section_content.extend([s.strip() for s in skills if s.strip() and len(s.strip()) > 1])
                elif current_section in ["education", "experience", "projects"]:
                    # Collect multi-line entries
                    if line_stripped and not re.match(r"^[A-Z\s]+$", line_stripped[:20]):  # Not just a header
                        section_content.append(line_stripped)
                else:
                    section_content.append(line_stripped)
        
        # Add remaining content
        if section_content and current_section:
            parsed[current_section].extend(section_content)
        
        # Post-process sections
        parsed["technical_skills"] = list(set([s for s in parsed["technical_skills"] if len(s) > 1]))[:50]
        
        # Clean up experience and projects (remove very short entries)
        parsed["experience"] = [e for e in parsed["experience"] if len(e) > 10]
        parsed["projects"] = [p for p in parsed["projects"] if len(p) > 10]
        
        return parsed
